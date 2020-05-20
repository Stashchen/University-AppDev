import sys

import docx
from PIL import Image

import PyQt5.QtWidgets as QtWidgets
from forms.design import Ui_MainWindow

from entity.fucn_1 import FirstDiff
from entity.func_2 import SecondDiff
from entity.func_3 import ThirdDiff


class MainMenu(QtWidgets.QMainWindow, Ui_MainWindow):

    def __init__(self):
        super().__init__()
        self.setupUi(self)

        self.graphicsView.setBackground('w')
        self.graphicsView.showGrid(x=True, y=True)

        self.btn.clicked.connect(self.draw)
        self.word_btn.clicked.connect(self.get_result_in_word)

    def draw(self):
        self.graphicsView.clear()

        if self.func_1.isChecked():
            self.equation = FirstDiff(float(self.start_edit.text()),
                                      float(self.end_edit.text()),
                                      float(self.accuracy_edit.text())
                                      )
        elif self.func_2.isChecked():
            self.equation = SecondDiff(float(self.start_edit.text()),
                                       float(self.end_edit.text()),
                                       float(self.accuracy_edit.text())
                                       )
        else:
            self.equation = ThirdDiff(float(self.start_edit.text()),
                                      float(self.end_edit.text()),
                                      float(self.accuracy_edit.text())
                                      )

        self.equation.reset_x_y()

        if self.euler_method.isChecked():
            self.equation.count_Euler()
        elif self.ext_euler_method.isChecked():
            self.equation.count_ext_Euler()
        elif self.runge_kut_3_method.isChecked():
            self.equation.count_runge_kut_third()
        else:
            self.equation.count_runge_kut_forth()

        self.equation.draw(self.graphicsView)
        self.change_table()

    def change_table(self):
        x_lst = self.equation.x_lst
        y_lst = self.equation.y_lst

        self.result_table.setColumnCount(len(x_lst))
        self.result_table.setHorizontalHeaderLabels(["N_{}".format(index) for index in range(1, 1 + len(x_lst))])

        for index, val in enumerate(x_lst):
            self.result_table.setItem(0, index, QtWidgets.QTableWidgetItem('{:.2f}'.format(val)))
            self.result_table.setItem(1, index, QtWidgets.QTableWidgetItem('{:.4f}'.format(y_lst[index])))

    def get_result_in_word(self):
        import pyautogui
        image = pyautogui.screenshot()
        image.save('./additional/image.jpeg')
        image = Image.open('./additional/image.jpeg')
        cropped = image.crop((240, 100, 980, 530))
        cropped.save('./additional/cropped.jpeg')

        docs = docx.Document()
        parag = docs.add_paragraph()
        run = parag.add_run()
        run.add_text('Результат')
        parag = docs.add_paragraph()
        run = parag.add_run()
        run.add_picture('./additional/cropped.jpeg', width=docx.shared.Inches(6), height=docx.shared.Inches(5))
        docs.save('./additional/word.docx')

if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    window = MainMenu()
    window.show()
    app.exec_()
